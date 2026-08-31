from __future__ import annotations

from pathlib import Path

from app.extractors.base import ExtractedDocument, ExtractedChunk

import logging

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extract text from DOCX files using python-docx."""

    SUPPORTED = {".docx"}

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() in self.SUPPORTED

    def extract(self, path: Path) -> ExtractedDocument:
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed. Cannot extract: %s", path)
            return ExtractedDocument(
                chunks=[],
                metadata={"format": "docx", "error": "python-docx not installed"},
            )

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        chunks: list[ExtractedChunk] = []

        # Try to split by headings (paragraphs with heading style)
        current_section = ""
        section_lines: list[str] = []
        offset = 0

        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                # Save previous section
                if section_lines:
                    section_text = "\n".join(section_lines)
                    chunks.append(ExtractedChunk(
                        text=section_text,
                        start_offset=offset - len(section_text),
                        end_offset=offset,
                        section=current_section,
                    ))

                current_section = para.text.strip()
                section_lines = [para.text]
            else:
                if para.text.strip():
                    section_lines.append(para.text)

            offset += len(para.text) + 2  # +2 for \n\n

        # Last section
        if section_lines:
            section_text = "\n".join(section_lines)
            chunks.append(ExtractedChunk(
                text=section_text,
                start_offset=max(0, offset - len(section_text)),
                end_offset=offset,
                section=current_section,
            ))

        # Fallback: if no headings found, use full text
        if not chunks and full_text:
            chunks.append(ExtractedChunk(
                text=full_text,
                start_offset=0,
                end_offset=len(full_text),
            ))

        return ExtractedDocument(
            chunks=chunks,
            metadata={"format": "docx"},
        )
